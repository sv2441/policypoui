import streamlit as st
import pandas as pd
import base64
import csv
import torch
import math
import docx
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()
os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]

from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import ResponseSchema
from langchain.output_parsers import StructuredOutputParser
from langchain.prompts import ChatPromptTemplate

# Initialize chat model
chat_llm = ChatOpenAI(temperature=0.0)

# Function to convert dictionary to CSV
def dict_to_csv(data, filename, append=False):
    mode = 'a' if append else 'w'
    with open(filename, mode, newline='') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=data.keys())
        if not append:
            writer.writeheader()
        writer.writerow(data)

def combine_policy(df):
    grouped_df = df.groupby('Article')['Policy'].apply(lambda x: ' '.join(x)).reset_index()
    grouped_df.rename(columns={'Policy': 'Policies'}, inplace=True)
    grouped_df.to_csv('grouped.csv', index=False)
    return grouped_df


def combine_column_to_paragraph(df, column_name):
    column_data = df[column_name].tolist()
    paragraph = " ".join(str(item) for item in column_data)
    return paragraph



# Function to process the CSV and perform policy generation
def result(df):
    Policy_schema = ResponseSchema(name="Policy", description="Policy Statement")
    response_schemas = [Policy_schema]
    output_parser = StructuredOutputParser.from_response_schemas(response_schemas)
    format_instructions = output_parser.get_format_instructions()
    
    title_template = """ \ You are an AI Governance bot. Just Execute the set of steps one by one.
                Convert "{topic}" into policy statements in 10 words.
                {format_instructions}
                """ 
    prompt = ChatPromptTemplate.from_template(template=title_template)
    df2 = df["Combined"]
    
    for i in range(len(df2)):
        messages = prompt.format_messages(topic=df2[i], format_instructions=format_instructions)
        response = chat_llm(messages)
        response_as_dict = output_parser.parse(response.content)
        data = response_as_dict
        dict_to_csv(data, 'policy.csv', append=True)
    
    result = pd.read_csv("policy.csv", names=['Policy'])
    final = pd.concat([df, result], axis=1)
    final.to_csv("result1.csv")
    st.dataframe(final)

# Function to perform paraphrasing
# def result2(df):
#     df4 = df['Policy']
#     paraphrase = []
#     for i in df['Policy']:
#         a = get_response(i, 1)
#         paraphrase.append(a)

#     df['paraphrased_text'] = paraphrase
#     df.to_csv('result2.csv', index=False)
#     st.markdown("Paraphrased Successfully")

# # Function to perform topic generation and summarization
# def result3(df):
    

#     df5 = df['Policy'].str.replace('[\[\]]', '', regex=True)
#     # df['Policy'] = df['Policy'].str.replace('[\[\]]', '', regex=True)

#     # Combine all rows into a single variable
#     # combined_text = ' '.join(df['Policy'].astype(str))
    
#     title_template = """ \ You are an AI Governance bot. 
#                 summarize the "{topic}" as brief instructional policy statements with "{article}" as title.
#                 """ 
#     # for 10 batch classify and group the statements into key topic
#     prompt = ChatPromptTemplate.from_template(template=title_template)
    
    
#     if os.path.exists('test.doc'):
#             doc = docx.Document('test.doc')
#     else:
#             doc = docx.Document()
            
#     batches = group_into_batches(df)        
#     for name, data in batches:
#         paragraph = " ".join(data['Policy'])
#         messages = prompt.format_messages(topic=paragraph , article=name)
#         response = chat_llm(messages)
#         content = str(response.content)  # Assuming response.content is a string
#         doc.add_paragraph(content)

#     doc.save('test.doc')
#         # batch_size =5
#     # batches = split_into_batches(combined_text, batch_size)

#     # if os.path.exists('test.doc'):
#     #     doc = docx.Document('test.doc')
#     # else:
#     #     doc = docx.Document()

#     # for batch in batches:
#     #     paragraph = ". ".join(batch)
#     #     messages = prompt.format_messages(topic=paragraph)
#     #     response = chat_llm(messages)
#     #     content = str(response.content)  # Assuming response.content is a string
#     #     doc.add_paragraph(content)

#     # doc.save('test.doc')
#     # doc = docx.Document() 
#     # messages = prompt.format_messages(topic=combined_text)
#     # response = chat_llm(messages)
#     # content = str(response.content)  # Assuming response.content is a string
#     # doc.add_paragraph(content)
#     # doc.save('test.doc')

#     with open('test.doc', 'rb') as f:
#         doc_data = f.read()
#     b64 = base64.b64encode(doc_data).decode()
#     href = f'<a href="data:application/octet-stream;base64,{b64}" download="result.doc">Download Result</a>'
#     st.markdown(href, unsafe_allow_html=True)
    
    
# # Function to concatenate columns and download the modified CSV file
def process_csv(df):
    # df = pd.read_csv(file)
    df['Combined'] = df['OP Title'].astype(str) + df['OP Description'].astype(str)
    modified_csv = df.to_csv(index=False)
    
    # Download the modified CSV file
    b64 = base64.b64encode(modified_csv.encode()).decode()
    href = f'<a href="data:file/csv;base64,{b64}" download="modified.csv">Download modified CSV file</a>'
    st.markdown(href, unsafe_allow_html=True)
    
    # Show a preview of the modified data
    st.dataframe(df)
    df.to_csv('policy_md1.csv')
    
def result2(df):
    combine_policy(df) 
    new_df=pd.read_csv('grouped.csv')
    
    Summary_schema = ResponseSchema(name="Summary",
                              description="Summary as instructional policy paragraph ")

    response_schemas = [Summary_schema]

    output_parser = StructuredOutputParser.from_response_schemas(response_schemas)
    format_instructions = output_parser.get_format_instructions()
    
    title_template = """ \ You are an AI Governance bot.  
                    Summarize "{topic}" as instructional policy paragraph in {count} Words In Legal Language Style. 
                     {format_instructions}          
                """ 
    prompt = ChatPromptTemplate.from_template(template=title_template)
    
    for index, row in new_df.iterrows():
        topic = row['Policies']
        # Calculate 30% of the length of the 'topic' variable and store it in the variable 'count'
        word_count_topic = len(topic.split())
        count = math.ceil(0.2 * word_count_topic) 
        
        messages = prompt.format_messages(topic=topic, count=count,format_instructions = format_instructions)
        response = chat_llm(messages)
        content = str(response.content)
        response_as_dict = output_parser.parse(response.content)
        data = response_as_dict
        dict_to_csv(data, 'summary.csv', append=True)
    result=pd.read_csv("summary.csv" , names=['Summary'])
    final=pd.concat([new_df, result], axis=1)
    final.to_csv("result2.csv")
    st.dataframe(final)

def result3(df):
    
    # last_df=pd.read_csv("result2.csv")
    contents = combine_column_to_paragraph(df, 'Summary')
    
    title_template = """ \ You are an AI Governance bot.  
                    Restructure {topic} based on topic into a policy document.          
                """ 
    prompt = ChatPromptTemplate.from_template(template=title_template)
    
    
    if os.path.exists('Policy_Document.doc'):
            doc = docx.Document('Policy_Document.doc')
    else:
            doc = docx.Document()
            
    
    messages = prompt.format_messages(topic=contents)
    response = chat_llm(messages)
    content = str(response.content)  # Assuming response.content is a string
    doc.add_paragraph(content)
    doc.save('Policy_Document.doc')


    with open('Policy_Document.doc', 'rb') as f:
        doc_data = f.read()
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="result.doc">Download Result</a>'
    st.markdown(href, unsafe_allow_html=True)
    
# Streamlit app
def main():
    st.title("Policy Prodago")
    
    # File upload
    uploaded_file = st.file_uploader("Upload CSV file", type=["csv"])
    
    
    if uploaded_file is not None:
        st.write("File uploaded successfully!")
        df = pd.read_csv(uploaded_file)
        
        st.subheader("CSV File Preview")
        st.dataframe(df)
        # Process button
        if st.button("Process"):
            process_csv(df)
        
        if st.button("Policy Generation"):
            result(pd.read_csv("policy_md1.csv"))
            
        if st.button("Summary"):
            result2(pd.read_csv('result1.csv',usecols=['Policy','Article']))
            
        if st.button("Policy Document Generation"):
            result3(pd.read_csv("result2.csv"))

if __name__ == "__main__":
    main()
